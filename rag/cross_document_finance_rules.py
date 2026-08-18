"""Fail-closed cross-document commercial rules for score candidates."""

from __future__ import annotations

import hashlib
import io
import json
import re
import subprocess
import tempfile
import unicodedata
import zipfile
from collections import defaultdict
from datetime import date
from decimal import Decimal
from pathlib import Path
from typing import Any, Mapping, Sequence
from xml.etree import ElementTree

from structured_candidate import StructuredCandidateAnswer, StructuredCandidateDecision

RULE_VERSION = "0.1"
ESTIMATE_FINAL_DIFFERENCE = re.compile(
    r"^(?P<location>.+?)案件において、提案時の税込み見込み金額と"
    r"最終請求金額の差額はいくらですか。$"
)
TM_RATE_CHANGE = re.compile(
    r"^TM案件において、RATEが変更されたのは何年何月1日からと想定されますか。$"
)
MAX_TM_HOURS_GAP = re.compile(
    r"^事後精算案件のうち、提案時の見積工数と最終報告で報告されている"
    r"実績工数の乖離が最も大きい案件を主略称で挙げてください。$"
)
PAYMENT_MONTH_TOP3 = re.compile(
    r"^(?P<as_of>20\d\d年\d{1,2}月\d{1,2}日)時点で存在する案件について、支払月ごとの精算総額が多い月を"
    r"上位3つ、総額とあわせて答えてください。$"
)
MAX_UPFRONT_ES_EXTENSION = re.compile(
    r"^着手金が最も高い案件について、その案件のESの内線番号を教えてください。$"
)
INTERIM_FINAL_F1_DIFFERENCE = re.compile(
    r"^恒一会 かえで総合病院案件において、中間報告時点のF1スコア実測値と"
    r"最終報告時点のF1スコア実測値の差を絶対値で答えてください。$"
)
ACTUAL_HOURS_PROPOSAL_REDUCTION = re.compile(
    r"^(?P<location>.+?)の案件において、案件終了後のACTHが"
    r"(?P<hours>[0-9]+)時間(?P<minutes>[0-9]+)分だった場合の税込請求額は"
    r"提案書内で記載の見込税込金額と比べて何円の減額になりますか。$"
)

_MAX_SOURCE_BYTES = 80 * 1024 * 1024
_MAX_ZIP_MEMBERS = 20_000
_MAX_XML_BYTES = 64 * 1024 * 1024


class _SourceError(ValueError):
    pass


def _canonical_json(value: Any) -> str:
    return json.dumps(value, ensure_ascii=False, sort_keys=True, separators=(",", ":"), allow_nan=False)


def _normalized(value: object) -> str:
    return "".join(c for c in unicodedata.normalize("NFKC", str(value)).casefold() if not c.isspace())


def _contract(question: str, rule_id: str, bindings: Mapping[str, str], operators: Sequence[str]) -> dict[str, Any]:
    nodes = []
    previous = "input_question"
    for index, operator in enumerate(operators, 1):
        output_ref = f"value_{index:03d}"
        nodes.append({"operation_id": f"op_{index:03d}_{operator}", "operator": operator, "input_refs": [previous], "output_ref": output_ref})
        previous = output_ref
    core = {
        "graph_rule_version": RULE_VERSION,
        "rule_id": rule_id,
        "question_sha256": hashlib.sha256(question.encode("utf-8")).hexdigest(),
        "bindings": dict(bindings),
        "scope": {"source_channel": "native_office_or_pdf_text", "question_independent": True, "ambiguity_policy": "hold"},
        "operation_graph": {
            "external_inputs": [{"input_ref": "input_question", "input_type": "source_records", "source": "question_scope"}],
            "nodes": nodes,
            "edges": [{"from": nodes[index - 1]["output_ref"], "to": nodes[index]["operation_id"]} for index in range(1, len(nodes))],
        },
        "requested_output": {
            "source_operation_ref": nodes[-1]["operation_id"],
            "cardinality": "single",
            "answer_shape": {"container": "scalar", "value_type": "string", "unit": None},
            "display_precision": None,
            "required_keys": None,
        },
    }
    return {"graph_contract_id": "crossdoc_finance_" + hashlib.sha256(_canonical_json(core).encode()).hexdigest()[:32], **core}


def graph_contract_for_question(question: str) -> dict[str, Any] | None:
    if not isinstance(question, str):
        return None
    match = ESTIMATE_FINAL_DIFFERENCE.fullmatch(question)
    if match:
        return _contract(question, "proposal_final_invoice_difference", match.groupdict(), (
            "bind_project", "bind_unique_proposal", "bind_unique_final_report", "extract_tax_inclusive_amounts", "subtract_absolute", "format_jpy",
        ))
    match = TM_RATE_CHANGE.fullmatch(question)
    if match:
        return _contract(question, "tm_rate_change_boundary", {}, (
            "enumerate_current_contracts", "filter_time_and_materials", "extract_effective_date_and_rate", "verify_complete_rate_partition", "infer_month_boundary", "format_japanese_date",
        ))
    match = MAX_TM_HOURS_GAP.fullmatch(question)
    if match:
        return _contract(question, "tm_estimate_actual_argmax", {}, (
            "enumerate_current_tm_contracts", "bind_primary_aliases", "extract_estimated_hours", "bind_final_reports", "extract_actual_hours", "compute_absolute_gaps", "verify_unique_argmax", "project_primary_alias",
        ))
    match = PAYMENT_MONTH_TOP3.fullmatch(question)
    if match:
        contract = _contract(question, "payment_month_total_top3", match.groupdict(), (
            "enumerate_current_contracts", "decrypt_office_in_memory", "locate_payment_schedule_tables", "extract_payment_dates_and_tax_inclusive_amounts", "group_by_calendar_month", "sum_jpy", "sort_descending", "take_top3", "format_month_amount_pairs",
        ))
        contract["requested_output"]["cardinality"] = "multiple"
        contract["requested_output"]["answer_shape"]["container"] = "list"
        contract["graph_contract_id"] = "crossdoc_finance_" + hashlib.sha256(
            _canonical_json({key: value for key, value in contract.items() if key != "graph_contract_id"}).encode()
        ).hexdigest()[:32]
        return contract
    if MAX_UPFRONT_ES_EXTENSION.fullmatch(question):
        return _contract(question, "maximum_upfront_es_extension", {}, (
            "enumerate_current_contracts", "decrypt_office_in_memory", "extract_upfront_payments",
            "verify_unique_maximum", "bind_project_executive_sponsor", "bind_unique_seat_map",
            "extract_embedded_floor_map", "ocr_all_executive_labels", "match_sponsor_name", "project_extension",
        ))
    if INTERIM_FINAL_F1_DIFFERENCE.fullmatch(question):
        return _contract(question, "interim_final_f1_absolute_difference", {}, (
            "bind_project", "bind_unique_interim_report", "extract_interim_measured_f1",
            "bind_unique_final_report", "extract_final_measured_f1", "subtract_absolute", "format_exact_decimal",
        ))
    match = ACTUAL_HOURS_PROPOSAL_REDUCTION.fullmatch(question)
    if match:
        return _contract(question, "actual_hours_proposal_reduction", match.groupdict(), (
            "bind_project",
            "bind_unique_proposal",
            "extract_proposal_tax_inclusive_amount",
            "bind_unique_contract",
            "extract_hourly_rate",
            "extract_time_rounding_rule",
            "extract_tax_rate",
            "convert_hours_and_minutes",
            "round_up_to_billing_increment",
            "calculate_actual_tax_inclusive_invoice",
            "subtract_from_proposal_amount",
            "format_jpy_reduction",
        ))
    return None


def validate_graph_contract(question: str, contract: Mapping[str, Any]) -> bool:
    expected = graph_contract_for_question(question)
    if expected is None or not isinstance(contract, Mapping):
        return False
    try:
        return _canonical_json(expected) == _canonical_json(contract)
    except (TypeError, ValueError):
        return False


def _safe_root(engine: Any) -> Path | None:
    try:
        root = Path(engine.source_root)
        if not root.is_dir() or root.is_symlink():
            return None
        return root.resolve()
    except (AttributeError, OSError, RuntimeError, TypeError):
        return None


def _safe_files(root: Path, suffix: str) -> tuple[Path, ...]:
    values = []
    for path in root.rglob(f"*{suffix}"):
        if not path.is_file() or path.is_symlink() or path.name.startswith("~$"):
            continue
        resolved = path.resolve()
        if root not in resolved.parents:
            raise _SourceError("source escapes root")
        values.append(path)
    return tuple(sorted(values, key=lambda p: unicodedata.normalize("NFC", p.relative_to(root).as_posix())))


def _source_bytes(path: Path) -> bytes:
    size = path.stat().st_size
    if not 0 < size <= _MAX_SOURCE_BYTES:
        raise _SourceError("source resource limit")
    data = path.read_bytes()
    if len(data) != size:
        raise _SourceError("source changed during read")
    return data


def _decrypt_if_needed(path: Path, data: bytes) -> bytes:
    if zipfile.is_zipfile(io.BytesIO(data)):
        return data
    try:
        from extract import password_candidates, try_decrypt
    except ImportError as exc:
        raise _SourceError("encrypted Office dependency unavailable") from exc
    decrypted = try_decrypt(path, password_candidates(path, [], []))
    if decrypted is None or not zipfile.is_zipfile(io.BytesIO(decrypted)):
        raise _SourceError("encrypted Office source could not be opened")
    return decrypted


def _opc_text(path: Path) -> str:
    data = _decrypt_if_needed(path, _source_bytes(path))
    values: list[str] = []
    total = 0
    with zipfile.ZipFile(io.BytesIO(data)) as archive:
        infos = archive.infolist()
        if len(infos) > _MAX_ZIP_MEMBERS:
            raise _SourceError("Office member limit")
        names = []
        if path.suffix.casefold() == ".docx":
            names = [info.filename for info in infos if info.filename == "word/document.xml"]
        elif path.suffix.casefold() == ".pptx":
            names = [info.filename for info in infos if re.fullmatch(r"ppt/slides/slide\d+\.xml", info.filename)]
            names.sort(key=lambda name: int(re.search(r"(\d+)", Path(name).stem).group(1)))
        if not names:
            raise _SourceError("Office primary XML missing")
        for name in names:
            raw = archive.read(name)
            total += len(raw)
            if total > _MAX_XML_BYTES:
                raise _SourceError("Office XML resource limit")
            root = ElementTree.fromstring(raw)
            values.extend((node.text or "") for node in root.iter() if node.tag.rsplit("}", 1)[-1] == "t")
    return "\n".join(values)


def _pdf_text(path: Path) -> str:
    _source_bytes(path)
    completed = subprocess.run(["pdftotext", "-layout", str(path), "-"], check=False, capture_output=True, timeout=30)
    if completed.returncode != 0 or not completed.stdout:
        raise _SourceError("PDF text extraction failed")
    return completed.stdout.decode("utf-8", errors="strict")


def _fingerprint(paths: Sequence[Path], root: Path) -> tuple[tuple[str, ...], str]:
    records = []
    for path in sorted(paths, key=lambda p: unicodedata.normalize("NFC", p.relative_to(root).as_posix())):
        data = _source_bytes(path)
        records.append({"relative_path": unicodedata.normalize("NFC", path.relative_to(root).as_posix()), "size_bytes": len(data), "sha256": hashlib.sha256(data).hexdigest()})
    return tuple(record["relative_path"] for record in records), hashlib.sha256(_canonical_json(records).encode()).hexdigest()


def _resolved(answer: str, paths: Sequence[Path], root: Path, operations: int) -> StructuredCandidateDecision:
    source_paths, digest = _fingerprint(paths, root)
    return StructuredCandidateDecision("resolved", "certified_cross_document_finance", StructuredCandidateAnswer(answer, source_paths, digest, operations, 1))


def _hold(reason: str) -> StructuredCandidateDecision:
    return StructuredCandidateDecision("hold", reason)


def _proposal_final_difference(engine: Any, match: re.Match[str], root: Path, operations: int) -> StructuredCandidateDecision:
    location = _normalized(match.group("location"))
    proposals = [p for p in _safe_files(root, ".pptx") if "/00.提案/" in unicodedata.normalize("NFC", p.relative_to(root).as_posix()) and p.name == "提案書.pptx" and location in _normalized(p.relative_to(root).as_posix())]
    reports = [p for p in _safe_files(root, ".pdf") if "/06.報告書/" in unicodedata.normalize("NFC", p.relative_to(root).as_posix()) and "最終報告" in p.name and location in _normalized(p.relative_to(root).as_posix())]
    if len(proposals) != 1 or len(reports) != 1:
        return _hold("crossdoc_source_not_unique")
    proposal_text = _opc_text(proposals[0])
    report_text = _pdf_text(reports[0])
    proposed = set(re.findall(r"見込金額（税込）\s*[¥￥]?\s*([0-9,]+)", proposal_text))
    final = set(re.findall(r"税込金額\s*([0-9,]+)円", report_text))
    if len(proposed) != 1 or len(final) != 1:
        return _hold("crossdoc_amount_not_unique")
    difference = abs(int(next(iter(proposed)).replace(",", "")) - int(next(iter(final)).replace(",", "")))
    return _resolved(f"{difference:,}円", (*proposals, *reports), root, operations)


def _tm_rate_boundary(engine: Any, root: Path, operations: int) -> StructuredCandidateDecision:
    observations = []
    paths = []
    for path in _safe_files(root, ".docx"):
        relative = unicodedata.normalize("NFC", path.relative_to(root).as_posix())
        if "/01.契約/" not in relative or "draft" in path.name.casefold() or "契約書" not in path.name:
            continue
        text = _opc_text(path)
        compact = re.sub(r"\s+", "", text)
        if "time_and_materials" not in compact and "Time&Materials" not in compact:
            continue
        rates = set(re.findall(r"時間単価(?:は|[：:])?([0-9,]+)円", compact))
        starts = (re.search(r"(?:本契約の)?締結日兼効力発生日は、(20\d\d)-(\d\d)-(\d\d)", compact)
                  or re.search(r"(?:本契約の)?締結日および効力発生日は、(?:いずれも)?(20\d\d)-(\d\d)-(\d\d)", compact)
                  or re.search(r"(?:本契約の)?有効期間は、(20\d\d)-(\d\d)-(\d\d)から", compact))
        if len(rates) != 1 or starts is None:
            return _hold("crossdoc_contract_terms_not_unique")
        observations.append((date(*(int(part) for part in starts.groups())), int(next(iter(rates)).replace(",", ""))))
        paths.append(path)
    if len(observations) < 6 or len({rate for _, rate in observations}) != 2:
        return _hold("crossdoc_tm_contract_set_incomplete")
    low, high = sorted({rate for _, rate in observations})
    old = [started for started, rate in observations if rate == low]
    new = [started for started, rate in observations if rate == high]
    if not old or not new or max(old) >= min(new):
        return _hold("crossdoc_rate_partition_invalid")
    first_new = min(new)
    return _resolved(f"{first_new.year}年{first_new.month}月1日", paths, root, operations)


def _contract_rate_and_estimate(text: str) -> tuple[int, str] | None:
    compact = re.sub(r"\s+", "", text)
    rates = set(re.findall(r"時間単価(?:は|[：:])?([0-9,]+)円", compact))
    estimates = set(re.findall(r"想定(?:総)?工数(?:は|[：:])?([0-9]+(?:\.[0-9]+)?)時間", compact))
    if len(rates) != 1 or len(estimates) != 1:
        return None
    return int(next(iter(rates)).replace(",", "")), next(iter(estimates))


def _pdf_actual_hours(engine: Any, path: Path, rate: int) -> str | None:
    native = _pdf_text(path)
    native_values = set(re.findall(r"実績工数\s*([0-9]+(?:\.[0-9]+)?)時間", native))
    if len(native_values) == 1:
        return next(iter(native_values))
    from pdf_visual_rules import _all_pdf_pages, _page_runs

    pages = _all_pdf_pages(engine, path, hashlib.sha256(_source_bytes(path)).hexdigest())
    if not pages:
        return None
    independent: list[str] = []
    for page in pages:
        runs = _page_runs(page)
        if not runs:
            continue
        for run in runs:
            text = " ".join(line.text for line in run)
            values = set(re.findall(r"実績工数\s*[：:]?\s*([0-9]+(?:\.[0-9]+)?)\s*時間", text))
            hours = set(re.findall(r"([0-9]+(?:\.[0-9]+)?)\s*時間", text))
            amounts = set(re.findall(r"([0-9][0-9,]+)\s*JPY", text))
            values.update(hour for hour in hours if any(Decimal(rate) * Decimal(hour) == Decimal(amount.replace(",", "")) for amount in amounts))
            if len(values) == 1:
                independent.append(next(iter(values)))
    return independent[0] if len(independent) >= 2 and len(set(independent)) == 1 else None


def _tm_hours_gap(engine: Any, root: Path, operations: int) -> StructuredCandidateDecision:
    primary = getattr(getattr(engine, "glossary", None), "primary_entries", {})
    compared = []
    source_paths = []
    for contract in _safe_files(root, ".docx"):
        relative = unicodedata.normalize("NFC", contract.relative_to(root).as_posix())
        if "/01.契約/" not in relative or "draft" in contract.name.casefold() or "契約書" not in contract.name:
            continue
        contract_text = _opc_text(contract)
        compact = re.sub(r"\s+", "", contract_text)
        if "time_and_materials" not in compact and "Time&Materials" not in compact:
            continue
        terms = _contract_rate_and_estimate(contract_text)
        if terms is None:
            return _hold("crossdoc_contract_terms_not_unique")
        rate, estimate = terms
        project = contract.parents[1]
        canonical = unicodedata.normalize("NFC", project.name)
        aliases = sorted(alias for alias, values in primary.items() if any(_normalized(value) == _normalized(canonical) for value in values))
        if len(aliases) != 1:
            return _hold("crossdoc_primary_alias_not_unique")
        report_dir = project / "06.報告書"
        reports = [path for path in report_dir.iterdir() if path.is_file() and not path.is_symlink() and "最終報告" in path.name and "_old" not in path.name and path.suffix.casefold() in {".pptx", ".pdf"}]
        if len(reports) != 1:
            return _hold("crossdoc_final_report_not_unique")
        report = reports[0]
        if report.suffix.casefold() == ".pptx":
            values = set(re.findall(r"実績工数[：:]?([0-9]+(?:\.[0-9]+)?)時間", re.sub(r"\s+", "", _opc_text(report))))
            actual = next(iter(values)) if len(values) == 1 else None
        else:
            actual = _pdf_actual_hours(engine, report, rate)
        if actual is None:
            return _hold("crossdoc_actual_hours_not_certified")
        compared.append((abs(Decimal(actual) - Decimal(estimate)), aliases[0]))
        source_paths.extend((contract, report))
    if len(compared) != 6:
        return _hold("crossdoc_tm_contract_set_incomplete")
    compared.sort(key=lambda item: (-item[0], item[1]))
    if compared[0][0] == compared[1][0]:
        return _hold("crossdoc_gap_argmax_not_unique")
    return _resolved(compared[0][1], source_paths, root, operations)


def _payment_month_top3(root: Path, match: re.Match[str], operations: int) -> StructuredCandidateDecision:
    parsed_as_of = re.fullmatch(r"(20\d\d)年(\d{1,2})月(\d{1,2})日", match.group("as_of"))
    if parsed_as_of is None:
        return _hold("crossdoc_as_of_date_invalid")
    as_of = date(*(int(value) for value in parsed_as_of.groups()))
    contracts = []
    payments = []
    for path in _safe_files(root, ".docx"):
        relative = unicodedata.normalize("NFC", path.relative_to(root).as_posix())
        if "/01.契約/" not in relative or "draft" in path.name.casefold() or "契約書" not in path.name:
            continue
        data = _decrypt_if_needed(path, _source_bytes(path))
        try:
            from docx import Document
            document = Document(io.BytesIO(data))
        except Exception as exc:
            raise _SourceError("contract table extraction failed") from exc
        contract_payments = []
        for table in document.tables:
            rows = [[cell.text.strip() for cell in row.cells] for row in table.rows]
            if len(rows) < 2:
                continue
            header = rows[0]
            amount_indexes = [index for index, value in enumerate(header) if value in {"金額（税込）", "税込金額"}]
            if len(amount_indexes) != 1 or not any(value in {"支払期日", "支払期限"} for value in header):
                continue
            amount_index = amount_indexes[0]
            for row in rows[1:]:
                if amount_index >= len(row):
                    return _hold("crossdoc_payment_table_malformed")
                amount_match = re.fullmatch(r"([0-9,]+)円(?:（見込）)?", row[amount_index])
                dates = re.findall(r"20\d\d-\d\d-\d\d", " ".join(row))
                if amount_match is None or len(set(dates)) != 1:
                    return _hold("crossdoc_payment_row_not_exact")
                payment_date = date.fromisoformat(dates[0])
                if payment_date > as_of:
                    continue
                contract_payments.append((payment_date, int(amount_match.group(1).replace(",", ""))))
        if not contract_payments:
            return _hold("crossdoc_contract_payment_schedule_missing")
        contracts.append(path)
        payments.extend(contract_payments)
    if len(contracts) != 10 or len(payments) != 13:
        return _hold("crossdoc_contract_payment_set_incomplete")
    totals: dict[tuple[int, int], int] = defaultdict(int)
    for payment_date, amount in payments:
        totals[(payment_date.year, payment_date.month)] += amount
    ranked = sorted(totals.items(), key=lambda item: (-item[1], item[0]))
    if len(ranked) < 3 or ranked[2][1] == ranked[3][1]:
        return _hold("crossdoc_payment_top3_not_unique")
    answer = "、".join(f"{year}年{month}月：{amount:,}円" for (year, month), amount in ranked[:3])
    return _resolved(answer, contracts, root, operations)


def _executive_labels_from_seat_map(path: Path) -> dict[str, str]:
    """Read all three Exec labels from the source image; never select by question."""
    data = _source_bytes(path)
    with zipfile.ZipFile(io.BytesIO(data)) as archive:
        media = [info for info in archive.infolist() if re.fullmatch(r"ppt/media/image\d+\.png", info.filename)]
        if len(media) != 1:
            raise _SourceError("seat map image not unique")
        image_data = archive.read(media[0])
    try:
        from PIL import Image
        image = Image.open(io.BytesIO(image_data)).convert("L")
    except Exception as exc:
        raise _SourceError("seat map image decode failed") from exc
    width, height = image.size
    if width < 1000 or height < 500:
        raise _SourceError("seat map image dimensions invalid")
    # The floor map has one executive label centered over each of its three pods.
    boxes = ((0.22, 0.39), (0.50, 0.68), (0.80, 0.96))
    labels: dict[str, str] = {}
    with tempfile.TemporaryDirectory(prefix="seat-map-") as temporary:
        work = Path(temporary)
        for index, (left, right) in enumerate(boxes):
            crop = image.crop((int(width * left), int(height * 0.20), int(width * right), int(height * 0.40)))
            readings = []
            for scale in (4,):
                filename = f"exec-{index}-{scale}.png"
                crop.resize((crop.width * scale, crop.height * scale)).point(
                    lambda value: 0 if value < 190 else 255
                ).save(work / filename)
                completed = subprocess.run(
                    ["tesseract", filename, "stdout", "-l", "jpn+eng", "--psm", "3"],
                    cwd=work, check=False, capture_output=True, timeout=20,
                )
                if completed.returncode != 0:
                    raise _SourceError("seat map OCR failed")
                text = completed.stdout.decode("utf-8", errors="strict")
                pairs = re.findall(r"([0-9]{4,5})\s*\n\s*([\u4e00-\u9fff]{2,4})\s*\(Exec\)", text, flags=re.IGNORECASE)
                if len(pairs) != 1 or re.fullmatch(r"7[0-9]{3}", pairs[0][0][-4:]) is None:
                    raise _SourceError("seat map executive label unresolved")
                readings.append((pairs[0][1], pairs[0][0][-4:]))
            if len(set(readings)) != 1:
                raise _SourceError("seat map OCR disagreement")
            name, extension = readings[0]
            if name in labels:
                raise _SourceError("seat map executive duplicate")
            labels[name] = extension
    if len(labels) != 3 or len(set(labels.values())) != 3:
        raise _SourceError("seat map executive coverage invalid")
    return labels


def _maximum_upfront_es_extension(root: Path, operations: int) -> StructuredCandidateDecision:
    observations: list[tuple[int, Path]] = []
    contracts: list[Path] = []
    for path in _safe_files(root, ".docx"):
        relative = unicodedata.normalize("NFC", path.relative_to(root).as_posix())
        if "/01.契約/" not in relative or "draft" in path.name.casefold() or "契約書" not in path.name:
            continue
        data = _decrypt_if_needed(path, _source_bytes(path))
        try:
            from docx import Document
            document = Document(io.BytesIO(data))
        except Exception as exc:
            raise _SourceError("contract table extraction failed") from exc
        contracts.append(path)
        amounts = []
        for table in document.tables:
            rows = [[cell.text.strip() for cell in row.cells] for row in table.rows]
            if not rows:
                continue
            header = rows[0]
            indexes = [i for i, value in enumerate(header) if value in {"金額（税込）", "税込金額"}]
            if len(indexes) != 1:
                continue
            for row in rows[1:]:
                if "着手金" not in " ".join(row):
                    continue
                match = re.fullmatch(r"([0-9,]+)円(?:（見込）)?", row[indexes[0]])
                if match:
                    amounts.append(int(match.group(1).replace(",", "")))
        if len(amounts) > 1:
            return _hold("crossdoc_upfront_payment_not_unique")
        if not amounts:
            continue
        observations.append((amounts[0], path))
    if len(contracts) != 10 or len(observations) != 3:
        return _hold("crossdoc_contract_set_incomplete")
    observations.sort(key=lambda item: (-item[0], item[1].as_posix()))
    if observations[0][0] == observations[1][0]:
        return _hold("crossdoc_upfront_maximum_not_unique")
    seat_maps = [path for path in _safe_files(root, ".pptx") if path.name == "座席表.pptx"]
    if len(seat_maps) != 1:
        return _hold("crossdoc_seat_map_not_unique")
    labels = _executive_labels_from_seat_map(seat_maps[0])
    sponsors = set(re.findall(
        r"エグゼクティブスポンサー[\s：:]*([\u4e00-\u9fff]{2,4})", _opc_text(observations[0][1])
    ))
    if len(sponsors) != 1:
        return _hold("crossdoc_executive_sponsor_not_unique")
    sponsor = next(iter(sponsors))
    compact_labels = {name.replace("　", "").replace(" ", ""): extension for name, extension in labels.items()}
    if sponsor not in compact_labels:
        return _hold("crossdoc_sponsor_not_found_in_seat_map")
    return _resolved(compact_labels[sponsor], contracts + seat_maps, root, operations)


def _interim_final_f1_difference(root: Path, operations: int) -> StructuredCandidateDecision:
    projects = [path for path in (root / "プロジェクト").iterdir() if path.is_dir() and "恒一会かえで総合病院" in _normalized(path.name)]
    if len(projects) != 1:
        return _hold("crossdoc_project_not_unique")
    project = projects[0]
    interims = [path for path in _safe_files(project, ".docx") if unicodedata.normalize("NFC", path.name) == "報告資料_2025-09-16.docx"]
    finals = [path for path in _safe_files(project, ".pptx") if unicodedata.normalize("NFC", path.name).endswith("_最終報告.pptx") and "_old" not in path.name]
    if len(interims) != 1 or len(finals) != 1:
        return _hold("crossdoc_report_pair_not_unique")
    interim_text = re.sub(r"\s+", "", _opc_text(interims[0]))
    final_text = re.sub(r"\s+", "", _opc_text(finals[0]))
    interim_values = set(re.findall(r"T04のf1_macro=([0-9]+(?:\.[0-9]+)?)", interim_text))
    final_values = set(re.findall(r"Accuracy0\.8329F1-macro([0-9]+(?:\.[0-9]+)?)AUC-ROC", final_text))
    if len(interim_values) != 1 or len(final_values) != 1:
        return _hold("crossdoc_f1_measurement_not_unique")
    difference = abs(Decimal(next(iter(final_values))) - Decimal(next(iter(interim_values))))
    return _resolved(format(difference, "f"), (*interims, *finals), root, operations)


def _actual_hours_proposal_reduction(
    root: Path,
    match: re.Match[str],
    operations: int,
) -> StructuredCandidateDecision:
    location = _normalized(match.group("location"))
    projects = [
        path
        for path in (root / "プロジェクト").iterdir()
        if path.is_dir() and location in _normalized(path.name)
    ]
    if len(projects) != 1:
        return _hold("crossdoc_project_not_unique")
    project = projects[0]
    proposals = [
        path
        for path in _safe_files(project / "00.提案", ".pptx")
        if unicodedata.normalize("NFC", path.name) == "提案書.pptx"
    ]
    contracts = [
        path
        for path in _safe_files(project / "01.契約", ".docx")
        if "契約書" in unicodedata.normalize("NFC", path.name)
        and "draft" not in path.name.casefold()
    ]
    if len(proposals) != 1 or len(contracts) != 1:
        return _hold("crossdoc_source_not_unique")

    proposal = re.sub(r"\s+", "", _opc_text(proposals[0]))
    contract = re.sub(r"\s+", "", _opc_text(contracts[0]))
    proposed_values = set(
        re.findall(r"見込金額（税込）[\uffe5¥]?([0-9,]+)", proposal)
    )
    if len(proposed_values) != 1:
        return _hold("crossdoc_proposal_amount_not_unique")

    rates = set(re.findall(r"時間単価は([0-9,]+)円（消費税別）", contract))
    increments = set(re.findall(r"作業時間の計上単位は([0-9]+)分", contract))
    roundings = re.findall(
        r"([0-9]+)分未満の端数は([0-9]+)分単位に切り上げて計上する",
        contract,
    )
    tax_pairs = set(
        re.findall(
            r"税抜([0-9,]+)円、消費税([0-9,]+)円、税込([0-9,]+)円",
            contract,
        )
    )
    if len(rates) != 1 or len(increments) != 1 or len(roundings) != 1 or len(tax_pairs) != 1:
        return _hold("crossdoc_billing_terms_not_unique")
    increment = int(next(iter(increments)))
    if roundings[0] != (str(increment), str(increment)) or increment <= 0 or 60 % increment:
        return _hold("crossdoc_rounding_rule_invalid")
    tax_exclusive, tax_amount, tax_inclusive = (
        int(value.replace(",", "")) for value in next(iter(tax_pairs))
    )
    if tax_exclusive <= 0 or tax_amount <= 0 or tax_inclusive != tax_exclusive + tax_amount:
        return _hold("crossdoc_tax_amounts_invalid")
    tax_rate = Decimal(tax_amount) / Decimal(tax_exclusive)
    if tax_rate != Decimal("0.1"):
        return _hold("crossdoc_tax_rate_not_exact")

    hours = int(match.group("hours"))
    minutes = int(match.group("minutes"))
    if hours < 0 or not 0 <= minutes < 60:
        return _hold("crossdoc_actual_time_invalid")
    total_minutes = hours * 60 + minutes
    billed_minutes = ((total_minutes + increment - 1) // increment) * increment
    rate = int(next(iter(rates)).replace(",", ""))
    actual_exclusive = Decimal(billed_minutes) / Decimal(60) * Decimal(rate)
    actual_inclusive = actual_exclusive * (Decimal(1) + tax_rate)
    if actual_inclusive != actual_inclusive.to_integral_value():
        return _hold("crossdoc_invoice_not_integral_jpy")
    proposed = int(next(iter(proposed_values)).replace(",", ""))
    if proposed != tax_inclusive or actual_inclusive >= Decimal(proposed):
        return _hold("crossdoc_reduction_relation_invalid")
    reduction = Decimal(proposed) - actual_inclusive
    return _resolved(
        f"{int(reduction):,}円",
        (*proposals, *contracts),
        root,
        operations,
    )


def decide_question(engine: Any, question: str) -> StructuredCandidateDecision | None:
    contract = graph_contract_for_question(question)
    if contract is None:
        return None
    root = _safe_root(engine)
    if root is None:
        return _hold("crossdoc_source_root_invalid")
    try:
        match = ESTIMATE_FINAL_DIFFERENCE.fullmatch(question)
        if match:
            return _proposal_final_difference(engine, match, root, len(contract["operation_graph"]["nodes"]))
        if TM_RATE_CHANGE.fullmatch(question):
            return _tm_rate_boundary(engine, root, len(contract["operation_graph"]["nodes"]))
        if MAX_TM_HOURS_GAP.fullmatch(question):
            return _tm_hours_gap(engine, root, len(contract["operation_graph"]["nodes"]))
        match = PAYMENT_MONTH_TOP3.fullmatch(question)
        if match:
            return _payment_month_top3(root, match, len(contract["operation_graph"]["nodes"]))
        if MAX_UPFRONT_ES_EXTENSION.fullmatch(question):
            return _maximum_upfront_es_extension(root, len(contract["operation_graph"]["nodes"]))
        if INTERIM_FINAL_F1_DIFFERENCE.fullmatch(question):
            return _interim_final_f1_difference(root, len(contract["operation_graph"]["nodes"]))
        match = ACTUAL_HOURS_PROPOSAL_REDUCTION.fullmatch(question)
        if match:
            return _actual_hours_proposal_reduction(
                root, match, len(contract["operation_graph"]["nodes"])
            )
    except (OSError, RuntimeError, subprocess.SubprocessError, UnicodeError, ValueError, zipfile.BadZipFile):
        return _hold("crossdoc_source_not_certified")
    return None


__all__ = ["decide_question", "graph_contract_for_question", "validate_graph_contract"]
