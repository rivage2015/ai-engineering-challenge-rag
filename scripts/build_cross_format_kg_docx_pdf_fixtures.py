#!/usr/bin/env python3
"""Build the DOCX and PDF fixtures for cross-format Knowledge Graph v0.1.

The two documents intentionally expose disjoint facts.  The DOCX defines the
project and work identities without people or assignment periods.  The PDF
defines employee-ID identity mappings without project or assignment facts.
Their separation prevents a single-document lookup from answering the later
cross-format evaluation questions.
"""

from __future__ import annotations

import argparse
import os
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor
from pypdf import PdfReader
from reportlab.lib.colors import HexColor
from reportlab.lib.pagesizes import A4
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFError, TTFont
from reportlab.pdfgen import canvas

if __package__:
    from .build_general_memory_docx_fixtures import (
        FIXED_TIME,
        configure as configure_base,
        save_deterministic,
        set_table_geometry,
        shade_cell,
    )
else:
    from build_general_memory_docx_fixtures import (
        FIXED_TIME,
        configure as configure_base,
        save_deterministic,
        set_table_geometry,
        shade_cell,
    )


PROJECT_DIRECTORY = "project-orion"
DOCX_NAME = "01_ORION-27_案件定義書.docx"
PDF_NAME = "05_社員ID対応表_署名済.pdf"

PDF_FONT_NAME = "CrossFormatJapanese"
PDF_FONT_CANDIDATES = (
    Path("/System/Library/Fonts/Supplemental/Arial Unicode.ttf"),
    Path("/System/Library/Fonts/Hiragino Sans GB.ttc"),
    Path("/System/Library/Fonts/ヒラギノ角ゴシック W3.ttc"),
    Path("/System/Library/Fonts/Hiragino Kaku Gothic ProN W3.otf"),
    Path("/Library/Fonts/Hiragino Sans GB.ttc"),
)
PDF_REQUIRED_TEXT = (
    "社員ID対応表",
    "EMP-104",
    "佐藤未来",
    "EMP-208",
    "高橋蓮",
    "2023-03-20",
    "署名済み",
    "承認済み",
)

PAGE_WIDTH, PAGE_HEIGHT = A4
NAVY = HexColor("#17324D")
BLUE = HexColor("#2E75B6")
PALE_BLUE = HexColor("#EAF4FB")
PALE_GRAY = HexColor("#F5F7F9")
INK = HexColor("#17212B")
MUTED = HexColor("#586875")
LINE = HexColor("#C9D5DF")
WHITE = HexColor("#FFFFFF")
DOCX_FONT = "Hiragino Kaku Gothic ProN"


def _configure_docx_run_properties(r_pr) -> None:
    r_fonts = r_pr.get_or_add_rFonts()
    for attr in ("asciiTheme", "hAnsiTheme", "eastAsiaTheme", "cstheme"):
        key = qn(f"w:{attr}")
        if key in r_fonts.attrib:
            del r_fonts.attrib[key]
    for attr in ("ascii", "hAnsi", "eastAsia", "cs"):
        r_fonts.set(qn(f"w:{attr}"), DOCX_FONT)
    lang = r_pr.find(qn("w:lang"))
    if lang is None:
        lang = OxmlElement("w:lang")
        r_pr.append(lang)
    lang.set(qn("w:val"), "ja-JP")
    lang.set(qn("w:eastAsia"), "ja-JP")


def set_font(run, size: float, bold: bool = False, color: str = "000000") -> None:
    run.font.name = DOCX_FONT
    _configure_docx_run_properties(run._element.get_or_add_rPr())
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def _configure_docx(document: Document, running_label: str) -> None:
    configure_base(document, running_label)
    doc_defaults = document.styles.element.find(qn("w:docDefaults"))
    r_pr_default = doc_defaults.find(qn("w:rPrDefault"))
    _configure_docx_run_properties(r_pr_default.find(qn("w:rPr")))
    for style_name in ("Normal", "Heading 1", "Heading 2", "Heading 3"):
        style = document.styles[style_name]
        style.font.name = DOCX_FONT
        _configure_docx_run_properties(style._element.get_or_add_rPr())


def _add_title(document: Document, title: str, subtitle: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(14)
    paragraph.paragraph_format.space_after = Pt(4)
    set_font(paragraph.add_run(title), 23, bold=True, color="0B2545")
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(16)
    set_font(paragraph.add_run(subtitle), 12, color="555555")


def _set_docx_cell_text(
    cell,
    value: str,
    *,
    bold: bool = False,
    centered: bool = False,
    color: str = "17212B",
) -> None:
    paragraph = cell.paragraphs[0]
    paragraph.alignment = (
        WD_ALIGN_PARAGRAPH.CENTER if centered else WD_ALIGN_PARAGRAPH.LEFT
    )
    paragraph.paragraph_format.space_after = Pt(0)
    set_font(paragraph.add_run(value), 10.5, bold=bold, color=color)


def build_project_definition_docx(path: Path) -> None:
    """Write identifiers only; people and assignment periods belong elsewhere."""
    document = Document()
    _configure_docx(document, "ORION-27 / Project Definition")
    document.core_properties.title = "ORION-27 案件定義書"
    document.core_properties.subject = "Cross-format Knowledge Graph evaluation fixture"
    document.core_properties.category = "FINAL project definition"
    document.core_properties.keywords = "ORION-27, WS-MIG-04, Project Orion"

    _add_title(document, "案件定義書", "ORION-27 / FINAL")
    introduction = document.add_paragraph(
        "本書は、案件と業務を形式横断で識別するための正式IDと別表記を定義する。"
    )
    introduction.paragraph_format.space_after = Pt(10)

    document.add_heading("識別情報", level=1)
    table = document.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    for index, value in enumerate(("項目", "定義値")):
        cell = table.rows[0].cells[index]
        shade_cell(cell, "DCEAF5")
        _set_docx_cell_text(
            cell, value, bold=True, centered=True, color="1F3A5F"
        )

    definitions = (
        ("Project ID", "ORION-27"),
        ("正式名称", "オリオン顧客移行"),
        ("Alias 1", "Project Orion"),
        ("Alias 2", "オリオン移行"),
        ("Work ID", "WS-MIG-04"),
        ("業務名", "移行リハーサル統括"),
        ("Status", "FINAL"),
    )
    for row_number, values in enumerate(definitions, start=1):
        cells = table.add_row().cells
        if row_number % 2 == 0:
            for cell in cells:
                shade_cell(cell, "F7F9FB")
        _set_docx_cell_text(cells[0], values[0], bold=True, color="1F3A5F")
        _set_docx_cell_text(cells[1], values[1])
    set_table_geometry(table, [2600, 6760])

    document.add_heading("同一性の扱い", level=1)
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(4)
    set_font(paragraph.add_run("案件別表記: "), 11, bold=True, color="2E75B6")
    set_font(
        paragraph.add_run(
            "Project Orionとオリオン移行は、Project ID ORION-27の別表記である。"
        ),
        11,
    )
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(4)
    set_font(paragraph.add_run("業務別表記: "), 11, bold=True, color="2E75B6")
    set_font(
        paragraph.add_run(
            "WS-MIG-04は、ORION-27に属する業務「移行リハーサル統括」の正式IDである。"
        ),
        11,
    )

    status = document.add_paragraph()
    status.paragraph_format.space_before = Pt(12)
    status.paragraph_format.space_after = Pt(0)
    set_font(status.add_run("文書状態: "), 10.5, bold=True, color="1F3A5F")
    set_font(status.add_run("FINAL"), 10.5, bold=True, color="177245")

    path.parent.mkdir(parents=True, exist_ok=True)
    save_deterministic(document, path)


def _register_embedded_japanese_font() -> Path:
    rejected: list[str] = []
    for candidate in PDF_FONT_CANDIDATES:
        if not candidate.is_file():
            continue
        try:
            pdfmetrics.registerFont(
                TTFont(PDF_FONT_NAME, str(candidate), subfontIndex=0)
            )
        except TTFError as error:
            rejected.append(f"{candidate}: {error}")
            continue
        return candidate
    searched = ", ".join(str(path) for path in PDF_FONT_CANDIDATES)
    detail = "; ".join(rejected) or "no candidate file existed"
    raise RuntimeError(
        "No embeddable Japanese font was found. Expected a Hiragino font at: "
        f"{searched}. Rejected candidates: {detail}"
    )


def _draw_pdf_text(
    pdf: canvas.Canvas,
    text: str,
    x: float,
    y: float,
    size: float,
    color=INK,
) -> None:
    pdf.setFillColor(color)
    pdf.setFont(PDF_FONT_NAME, size)
    pdf.drawString(x, y, text)


def _build_employee_identity_pdf(path: Path) -> None:
    pdf = canvas.Canvas(
        str(path), pagesize=A4, pageCompression=1, invariant=1
    )
    pdf.setTitle("社員ID対応表 / 2023-03-20")
    pdf.setAuthor("Cross-format Knowledge Graph Evaluation")
    pdf.setCreator("build_cross_format_kg_docx_pdf_fixtures.py")
    pdf.setSubject("署名済み・承認済みの社員ID対応表")

    pdf.setFillColor(BLUE)
    pdf.rect(0, PAGE_HEIGHT - 18, PAGE_WIDTH, 18, fill=1, stroke=0)
    _draw_pdf_text(pdf, "社員識別情報 / APPROVED REGISTER", 54, PAGE_HEIGHT - 54, 9, BLUE)
    _draw_pdf_text(pdf, "社員ID対応表", 54, PAGE_HEIGHT - 94, 23, NAVY)
    pdf.setStrokeColor(LINE)
    pdf.line(54, PAGE_HEIGHT - 112, PAGE_WIDTH - 54, PAGE_HEIGHT - 112)

    pdf.setFillColor(PALE_BLUE)
    pdf.roundRect(54, PAGE_HEIGHT - 186, PAGE_WIDTH - 108, 48, 7, fill=1, stroke=0)
    _draw_pdf_text(pdf, "名簿版: 2023-03-20", 72, PAGE_HEIGHT - 158, 11, INK)
    _draw_pdf_text(pdf, "文書状態: 署名済み/承認済み", 290, PAGE_HEIGHT - 158, 11, INK)

    left = 54
    top = PAGE_HEIGHT - 232
    column_widths = (150, 190, PAGE_WIDTH - 108 - 340)
    row_height = 42
    headers = ("社員ID", "氏名", "登録状態")
    rows = (
        ("EMP-104", "佐藤未来", "有効"),
        ("EMP-208", "高橋蓮", "有効"),
    )

    x = left
    for width, header in zip(column_widths, headers):
        pdf.setFillColor(NAVY)
        pdf.rect(x, top - row_height, width, row_height, fill=1, stroke=0)
        _draw_pdf_text(pdf, header, x + 12, top - 27, 10.5, WHITE)
        x += width

    for row_index, row in enumerate(rows, start=1):
        y = top - row_height * (row_index + 1)
        pdf.setFillColor(PALE_GRAY if row_index % 2 else WHITE)
        pdf.rect(left, y, sum(column_widths), row_height, fill=1, stroke=0)
        x = left
        for width, value in zip(column_widths, row):
            pdf.setStrokeColor(LINE)
            pdf.rect(x, y, width, row_height, fill=0, stroke=1)
            _draw_pdf_text(pdf, value, x + 12, y + 15, 11, INK)
            x += width

    approval_top = top - row_height * 4 - 34
    pdf.setFillColor(PALE_BLUE)
    pdf.roundRect(54, approval_top - 92, PAGE_WIDTH - 108, 92, 7, fill=1, stroke=0)
    _draw_pdf_text(pdf, "承認記録", 72, approval_top - 24, 12, NAVY)
    _draw_pdf_text(pdf, "承認主体: 人事情報管理部門", 72, approval_top - 49, 10.5, INK)
    _draw_pdf_text(pdf, "署名状態: 署名済み / 承認済み", 72, approval_top - 72, 10.5, INK)

    _draw_pdf_text(
        pdf,
        "合成評価用fixture - 外部データは含まない",
        54,
        32,
        8,
        MUTED,
    )
    _draw_pdf_text(pdf, "1 / 1", PAGE_WIDTH - 82, 32, 8, MUTED)
    pdf.save()


def _verify_pdf_text(path: Path) -> None:
    extracted = "\n".join(page.extract_text() or "" for page in PdfReader(path).pages)
    missing = [value for value in PDF_REQUIRED_TEXT if value not in extracted]
    if missing:
        raise RuntimeError(
            "Generated Japanese PDF is not fully extractable with pypdf; "
            f"missing={missing}"
        )


def build_employee_identity_pdf(path: Path) -> None:
    """Build atomically and reject a PDF without an extractable Unicode map."""
    _register_embedded_japanese_font()
    path.parent.mkdir(parents=True, exist_ok=True)
    temporary = path.with_name(f".{path.name}.building")
    try:
        _build_employee_identity_pdf(temporary)
        _verify_pdf_text(temporary)
        temporary.replace(path)
        timestamp = FIXED_TIME.timestamp()
        os.utime(path, (timestamp, timestamp))
    finally:
        temporary.unlink(missing_ok=True)


def main() -> int:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument(
        "--out",
        required=True,
        type=Path,
        help="Corpus root; fixtures are written below project-orion/",
    )
    parser.add_argument(
        "--only",
        choices=("docx", "pdf", "all"),
        default="all",
        help="Build one fixture format or both (default: all).",
    )
    args = parser.parse_args()
    project_root = args.out.resolve() / PROJECT_DIRECTORY
    if args.only in {"docx", "all"}:
        build_project_definition_docx(project_root / DOCX_NAME)
    if args.only in {"pdf", "all"}:
        build_employee_identity_pdf(project_root / PDF_NAME)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
